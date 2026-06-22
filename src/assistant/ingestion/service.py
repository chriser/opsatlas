"""Ingestion service: extract text from a registered source and build sections."""

from __future__ import annotations

from pathlib import Path

from ..sources.models import SourceRecord
from ..sources.register import SourceRegister
from .sections import build_sections
from .store import SectionStore

PLAIN_TEXT_EXTENSIONS = {".txt", ".md", ".json"}
SUPPORTED_EXTENSIONS = PLAIN_TEXT_EXTENSIONS | {".pdf", ".docx"}


class NotIngestableError(ValueError):
    """Raised when a source cannot be ingested (unsupported type, no text, etc.)."""


def extract_text(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension in PLAIN_TEXT_EXTENSIONS:
        return content.decode("utf-8", errors="replace")
    if extension == ".pdf":
        return _extract_pdf(content)
    if extension == ".docx":
        return _extract_docx(content)
    allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise NotIngestableError(f"Cannot extract text from '{extension or filename}'. Supported: {allowed}.")


def _extract_pdf(content: bytes) -> str:
    from io import BytesIO

    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(content))
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:  # pragma: no cover - corrupt file path
        raise NotIngestableError(f"Could not read the PDF: {exc}") from exc
    if not text.strip():
        raise NotIngestableError("No extractable text in the PDF (it may be scanned images).")
    return text


def _extract_docx(content: bytes) -> str:
    from io import BytesIO

    import docx

    try:
        document = docx.Document(BytesIO(content))
    except Exception as exc:
        raise NotIngestableError(f"Could not read the Word document: {exc}") from exc
    parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Promote Word headings to Markdown so the section builder splits on them.
        style = (para.style.name or "") if para.style else ""
        parts.append(f"# {text}" if style.startswith("Heading") else text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        raise NotIngestableError("No extractable text in the Word document.")
    return text


def ingest_source(
    register: SourceRegister,
    section_store: SectionStore,
    source_id: str,
) -> SourceRecord:
    record = register.get(source_id)
    if record is None:
        raise NotIngestableError("Source not found.")

    text = extract_text(record.filename, register.read_content(source_id))
    sections = build_sections(source_id, text)
    if not sections:
        section_store.remove_for_source(source_id)
        register.update(source_id, processing_state="failed", section_count=0)
        raise NotIngestableError(
            "No ingestible sections were found. Add body content below headings "
            "or include plain paragraphs, then ingest again."
        )

    section_store.replace_for_source(source_id, sections)

    updated = register.update(
        source_id, processing_state="ingested", section_count=len(sections)
    )
    assert updated is not None  # the record existed a moment ago
    return updated
