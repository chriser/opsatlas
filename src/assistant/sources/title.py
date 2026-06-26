"""Readable source title generation for uploaded documents."""

from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path
from typing import Any

MAX_TITLE_CHARS = 90

JSON_TITLE_KEYS = (
    "title",
    "document_title",
    "source_title",
    "process_title",
    "process_name",
    "name",
    "heading",
    "subject",
)

GENERIC_TITLES = {
    "document",
    "draft",
    "final",
    "introduction",
    "notes",
    "overview",
    "readme",
    "source",
    "title",
    "untitled",
}


def generate_source_title(filename: str, content: bytes, explicit_title: str | None = None) -> str:
    """Return a human-friendly title for the source register.

    Explicit titles are preserved. Otherwise we use document metadata or visible
    headings where possible, falling back to a cleaned-up filename.
    """

    explicit = _clean_candidate(explicit_title, preserve_case=True)
    if explicit:
        return explicit

    extension = Path(filename).suffix.lower()
    candidates: list[str | None] = []
    if extension in {".txt", ".md"}:
        candidates.append(_title_from_text(_decode_text(content)))
    elif extension == ".json":
        candidates.append(_title_from_json(content))
        candidates.append(_title_from_text(_decode_text(content)))
    elif extension == ".docx":
        candidates.extend(_title_from_docx(content))
    elif extension == ".pdf":
        candidates.extend(_title_from_pdf(content))

    for candidate in candidates:
        title = _clean_candidate(candidate, preserve_case=True)
        if _is_useful_title(title):
            return title

    return _title_from_filename(filename)


def _decode_text(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _title_from_json(content: bytes) -> str | None:
    try:
        parsed = json.loads(_decode_text(content))
    except json.JSONDecodeError:
        return None
    return _find_json_title(parsed)


def _find_json_title(value: Any) -> str | None:
    if isinstance(value, dict):
        lower_lookup = {str(key).lower(): key for key in value}
        for wanted in JSON_TITLE_KEYS:
            key = lower_lookup.get(wanted)
            if key is not None and isinstance(value[key], str):
                candidate = _clean_candidate(value[key], preserve_case=True)
                if _is_useful_title(candidate):
                    return candidate
        for nested in value.values():
            candidate = _find_json_title(nested)
            if candidate:
                return candidate
    elif isinstance(value, list):
        for item in value[:10]:
            candidate = _find_json_title(item)
            if candidate:
                return candidate
    return None


def _title_from_text(text: str) -> str | None:
    for raw_line in text.splitlines()[:60]:
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^#{1,3}\s+(.+)$", line)
        if heading:
            candidate = _clean_candidate(heading.group(1), preserve_case=True)
            if _is_useful_title(candidate):
                return candidate
            continue
        if line.startswith(("* ", "- ", "+ ")):
            continue
        if line.endswith((".", "!", "?")):
            continue
        candidate = _clean_candidate(line, preserve_case=True)
        if _looks_like_standalone_title(candidate):
            return candidate
    return None


def _title_from_docx(content: bytes) -> list[str | None]:
    try:
        import docx

        document = docx.Document(BytesIO(content))
    except Exception:
        return []

    candidates: list[str | None] = [document.core_properties.title]
    for para in document.paragraphs[:40]:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading"):
            candidates.append(text)
            break
    candidates.append(_title_from_text("\n".join(para.text for para in document.paragraphs[:40])))
    return candidates


def _title_from_pdf(content: bytes) -> list[str | None]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
    except Exception:
        return []

    metadata_title = getattr(reader.metadata, "title", None) if reader.metadata else None
    page_text: list[str] = []
    for page in reader.pages[:2]:
        try:
            page_text.append(page.extract_text() or "")
        except Exception:
            continue
    return [metadata_title, _title_from_text("\n".join(page_text))]


def _clean_candidate(candidate: str | None, *, preserve_case: bool) -> str:
    if not candidate:
        return ""
    title = candidate.strip()
    title = re.sub(r"^\ufeff", "", title)
    title = re.sub(r"^(title|document title|name|subject)\s*[:\-]\s*", "", title, flags=re.I)
    title = re.sub(r"^microsoft word\s*-\s*", "", title, flags=re.I)
    title = re.sub(r"^anonymised learning pack\s+\d+\s*[–-]\s*", "", title, flags=re.I)
    title = re.sub(r"\.(txt|md|pdf|docx|json)$", "", title, flags=re.I)
    title = re.sub(r"[_\t]+", " ", title)
    title = re.sub(r"\s+", " ", title).strip(" -_.:;")
    if not title:
        return ""
    if title.isupper() or not preserve_case:
        title = _smart_title_case(title)
    return _truncate_title(title)


def _title_from_filename(filename: str) -> str:
    stem = Path(filename).stem or "Uploaded Document"
    cleaned = re.sub(r"[_\-.]+", " ", stem)
    cleaned = re.sub(r"\b(copy|draft|final)\b", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\bv\d+\b", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\b\d{4}[-_ ]?\d{2}[-_ ]?\d{2}\b", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    title = _clean_candidate(cleaned, preserve_case=False)
    return title or "Uploaded Document"


def _looks_like_standalone_title(candidate: str) -> bool:
    if not _is_useful_title(candidate):
        return False
    words = candidate.split()
    if len(words) > 14:
        return False
    if candidate.endswith((".", "!", "?")):
        return False
    if len(candidate) > MAX_TITLE_CHARS:
        return False
    return True


def _is_useful_title(candidate: str) -> bool:
    if len(candidate) < 4:
        return False
    normalized = candidate.lower().strip()
    if normalized in GENERIC_TITLES:
        return False
    if re.fullmatch(r"[\d\W_]+", normalized):
        return False
    if len(candidate.split()) > 18:
        return False
    return True


def _smart_title_case(value: str) -> str:
    small_words = {"a", "an", "and", "as", "at", "for", "from", "in", "of", "on", "or", "the", "to", "with"}
    words = value.split()
    formatted: list[str] = []
    for index, word in enumerate(words):
        lower = word.lower()
        if index > 0 and lower in small_words:
            formatted.append(lower)
        elif word.isupper() and len(word) <= 5:
            formatted.append(word)
        else:
            formatted.append(lower[:1].upper() + lower[1:])
    return " ".join(formatted)


def _truncate_title(value: str) -> str:
    if len(value) <= MAX_TITLE_CHARS:
        return value
    trimmed = value[: MAX_TITLE_CHARS - 3].rsplit(" ", 1)[0].rstrip(" ,;:-")
    return f"{trimmed or value[: MAX_TITLE_CHARS - 3]}..."
