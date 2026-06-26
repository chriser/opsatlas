from assistant.sources.title import generate_source_title


def test_explicit_source_title_is_preserved():
    title = generate_source_title("supplier-process-v3-final.pdf", b"%PDF-1.4", "Supplier setup process")

    assert title == "Supplier setup process"


def test_source_title_uses_markdown_heading():
    title = generate_source_title(
        "supplier-process-v3-final.md",
        b"# Supplier onboarding controls\n\nCredit checks are mandatory before setup.",
    )

    assert title == "Supplier onboarding controls"


def test_source_title_uses_json_title_field():
    title = generate_source_title(
        "raw-export-001.json",
        b'{"metadata": {"document_title": "Article setup operating model"}, "items": []}',
    )

    assert title == "Article setup operating model"


def test_source_title_falls_back_to_clean_filename():
    title = generate_source_title(
        "supplier_process_v3_final_2026-06-26.txt",
        b"This paragraph is useful content but not a short standalone title.",
    )

    assert title == "Supplier Process"


def test_source_title_uses_docx_heading(tmp_path):
    import docx

    document = docx.Document()
    document.add_heading("Customer refund approval process", level=1)
    document.add_paragraph("Refund approvals require finance review.")
    path = tmp_path / "customer-refunds.docx"
    document.save(path)

    title = generate_source_title(path.name, path.read_bytes())

    assert title == "Customer refund approval process"
