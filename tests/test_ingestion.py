"""Tests for the section builder and the ingestion API."""

from fastapi.testclient import TestClient

from assistant.api.app import create_app
from assistant.api.auth import AuthService
from assistant.ingestion.sections import build_sections
from assistant.sources.register import SourceRegister

PASSWORD = "test-pass"

MARKDOWN = """# Supplier setup

Supplier setup begins with a business request.

# Controls

Due diligence and credit checks are mandatory gates.
"""


def make_client(tmp_path) -> TestClient:
    client = TestClient(create_app(SourceRegister(tmp_path), AuthService(PASSWORD)))
    token = client.post("/api/auth/login", json={"password": PASSWORD}).json()["token"]
    client.headers.update({"Authorization": f"Bearer {token}"})
    return client


def upload(client, name: str, body: bytes, content_type: str = "text/markdown"):
    return client.post(
        "/api/sources/upload",
        files={"file": (name, body, content_type)},
        data={"title": name},
    ).json()


# --- section builder (unit) ---

def test_build_sections_splits_on_headings():
    sections = build_sections("s1", MARKDOWN)
    assert [s.heading for s in sections] == ["Supplier setup", "Controls"]
    assert [s.ordinal for s in sections] == [0, 1]
    assert all(s.char_count == len(s.text) for s in sections)


def test_build_sections_handles_text_without_headings():
    sections = build_sections("s2", "Just one paragraph of plain text.")
    assert len(sections) == 1
    assert sections[0].heading == "Introduction"


# --- ingestion API ---

def test_ingest_advances_state_and_counts_sections(tmp_path):
    client = make_client(tmp_path)
    record = upload(client, "supplier.md", MARKDOWN.encode())
    assert record["processing_state"] == "registered"
    assert record["section_count"] == 0

    ingested = client.post(f"/api/sources/{record['id']}/ingest").json()
    assert ingested["processing_state"] == "ingested"
    assert ingested["section_count"] == 2

    sections = client.get(f"/api/sources/{record['id']}/sections").json()
    assert len(sections) == 2
    assert sections[0]["heading"] == "Supplier setup"


def test_ingest_rejects_unsupported_type(tmp_path):
    client = make_client(tmp_path)
    record = upload(client, "scan.pdf", b"%PDF-1.4 binary", "application/pdf")
    response = client.post(f"/api/sources/{record['id']}/ingest")
    assert response.status_code == 400


def test_ingest_missing_source_is_404(tmp_path):
    client = make_client(tmp_path)
    assert client.post("/api/sources/nope/ingest").status_code == 404


def test_ingest_requires_auth(tmp_path):
    client = make_client(tmp_path)
    record = upload(client, "supplier.md", MARKDOWN.encode())
    client.headers.pop("Authorization")
    assert client.post(f"/api/sources/{record['id']}/ingest").status_code == 401


# --- PDF / DOCX extraction ---

def test_extract_supports_json_as_text():
    from assistant.ingestion.service import extract_text
    assert "hello" in extract_text("a.json", b'{"note": "hello"}')


def test_extract_docx_promotes_headings_and_tables(tmp_path):
    import docx
    from assistant.ingestion.service import extract_text
    d = docx.Document()
    d.add_heading("Supplier setup", level=1)
    d.add_paragraph("Due diligence and credit checks are mandatory.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Role"
    t.rows[0].cells[1].text = "Requester"
    path = tmp_path / "x.docx"
    d.save(path)
    text = extract_text("x.docx", path.read_bytes())
    assert "# Supplier setup" in text
    assert "mandatory" in text
    assert "Role | Requester" in text


def test_extract_pdf_rejects_invalid_bytes():
    import pytest
    from assistant.ingestion.service import NotIngestableError, extract_text
    with pytest.raises(NotIngestableError):
        extract_text("x.pdf", b"not a real pdf")


def test_unsupported_extension_raises():
    import pytest
    from assistant.ingestion.service import NotIngestableError, extract_text
    with pytest.raises(NotIngestableError):
        extract_text("x.exe", b"\x00")
